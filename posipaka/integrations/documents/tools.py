"""Document generation integration — PDF, DOCX creation.

Generates documents from text/markdown content.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


async def generate_pdf(content: str, filename: str = "document.pdf", title: str = "") -> str:
    """Згенерувати PDF документ з тексту/markdown."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return "Для генерації PDF потрібен пакет reportlab. Встановіть: pip install reportlab"

    output_dir = Path.home() / ".posipaka" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if paragraph.startswith("# "):
            story.append(Paragraph(paragraph[2:], styles["Heading1"]))
        elif paragraph.startswith("## "):
            story.append(Paragraph(paragraph[3:], styles["Heading2"]))
        else:
            story.append(Paragraph(paragraph.replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))

    doc.build(story)
    return f"PDF створено: {output_path}"


async def generate_docx(content: str, filename: str = "document.docx", title: str = "") -> str:
    """Згенерувати DOCX документ з тексту/markdown."""
    try:
        from docx import Document
    except ImportError:
        return "Для генерації DOCX потрібен пакет python-docx. Встановіть: pip install python-docx"

    output_dir = Path.home() / ".posipaka" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    doc = Document()

    if title:
        doc.add_heading(title, 0)

    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if paragraph.startswith("# "):
            doc.add_heading(paragraph[2:], level=1)
        elif paragraph.startswith("## "):
            doc.add_heading(paragraph[3:], level=2)
        elif paragraph.startswith("### "):
            doc.add_heading(paragraph[4:], level=3)
        elif paragraph.startswith("- ") or paragraph.startswith("* "):
            for line in paragraph.split("\n"):
                line = line.lstrip("- *").strip()
                if line:
                    doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph(paragraph)

    doc.save(str(output_path))
    return f"DOCX створено: {output_path}"


async def generate_csv(headers: str, rows: str, filename: str = "data.csv") -> str:
    """Згенерувати CSV файл."""
    import csv

    output_dir = Path.home() / ".posipaka" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    header_list = [h.strip() for h in headers.split(",")]
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(header_list)
        for row in rows.strip().split("\n"):
            writer.writerow([c.strip() for c in row.split(",")])

    return f"CSV створено: {output_path}"


def register(registry: Any) -> None:
    from posipaka.core.tools.registry import ToolDefinition

    registry.register(
        ToolDefinition(
            name="generate_pdf",
            description="Generate a PDF document from text/markdown content.",
            category="documents",
            handler=generate_pdf,
            input_schema={
                "type": "object",
                "required": ["content"],
                "properties": {
                    "content": {"type": "string", "description": "Text or markdown content"},
                    "filename": {
                        "type": "string",
                        "description": "Output filename (default: document.pdf)",
                    },
                    "title": {"type": "string", "description": "Document title"},
                },
            },
            tags=["documents", "pdf", "generate"],
        )
    )

    registry.register(
        ToolDefinition(
            name="generate_docx",
            description="Generate a DOCX (Word) document from text/markdown content.",
            category="documents",
            handler=generate_docx,
            input_schema={
                "type": "object",
                "required": ["content"],
                "properties": {
                    "content": {"type": "string", "description": "Text or markdown content"},
                    "filename": {
                        "type": "string",
                        "description": "Output filename (default: document.docx)",
                    },
                    "title": {"type": "string", "description": "Document title"},
                },
            },
            tags=["documents", "docx", "generate"],
        )
    )

    registry.register(
        ToolDefinition(
            name="generate_csv",
            description="Generate a CSV file from headers and rows.",
            category="documents",
            handler=generate_csv,
            input_schema={
                "type": "object",
                "required": ["headers", "rows"],
                "properties": {
                    "headers": {"type": "string", "description": "Comma-separated headers"},
                    "rows": {
                        "type": "string",
                        "description": "Rows, one per line, comma-separated values",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Output filename (default: data.csv)",
                    },
                },
            },
            tags=["documents", "csv", "generate"],
        )
    )
